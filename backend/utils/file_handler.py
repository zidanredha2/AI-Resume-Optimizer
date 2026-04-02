import os
import tempfile
from werkzeug.utils import secure_filename
import docx

class FileHandler:
    """Handle file operations for resume uploads"""
    
    ALLOWED_EXTENSIONS = {'docx'}
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    
    @staticmethod
    def allowed_file(filename):
        """Check if file has allowed extension"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in FileHandler.ALLOWED_EXTENSIONS
    
    @staticmethod
    def validate_file_size(file):
        """Validate file size"""
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        return size <= FileHandler.MAX_FILE_SIZE
    
    @staticmethod
    def extract_text_from_docx(file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                text.append(para.text)
            
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    def save_uploaded_file(file):
        """Save uploaded file temporarily"""
        filename = secure_filename(file.filename)
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        return file_path
    
    @staticmethod
    def cleanup_file(file_path):
        """Remove temporary file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception:
            pass  # Silently fail on cleanup